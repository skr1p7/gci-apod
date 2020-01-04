from flask import Flask, render_template, request, redirect, url_for
import requests
import wget
from docx import Document
from docx.shared import Inches
import os
import convertapi

convertapi.api_secret = 'yPn1JnJhdFGsnvN8'

app = Flask(__name__)

@app.route('/', methods=['GET','POST'])
def main():
    if request.method == "POST":
        date = request.form['text']
        global url
        url = 'https://api.nasa.gov/planetary/apod?api_key=w43FMRhURuPXm5zO3mYsSa4pBJMfcH3CggwfyeZZ&date={}'.format(date)
        req = requests.get(url)
        data = req.json()
        if req.status_code == 400:
            return render_template('error.html')
        else:
            pass
        img = data['hdurl']
        title = data['title']
        return render_template('index.html',image=img, title=title)
    else:
        return render_template('base.html')

@app.route('/download/', methods=['GET','POST'])
def download():
    data = requests.get(url).json()
    imgUrl = data['hdurl']
    imgTitle = data['title']
    filename = wget.download(imgUrl)

    document = Document()
    document.add_heading(str(imgTitle), 0)
    document.add_picture(str(filename), width=Inches(6))
    document.add_page_break()
    document.save('file.docx')
    pdf = convertapi.convert('pdf', { 'File': 'file.docx' })
    pdfName = str(filename) + ".pdf"
    pdf.file.save(str(pdfName))
    os.system('rm -rf file.docx')
    removeImg = 'rm '+str(filename)
    os.system(str(removeImg))
    a = "File downloaded as " + filename

    
    return render_template('downloaded.html', a=a)

if __name__ == "__main__":
    app.run()